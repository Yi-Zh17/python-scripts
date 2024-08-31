# coding:utf-8
import sys
import os
from docx import Document
from datetime import date

default_encoding = "utf-8"
if default_encoding != sys.getdefaultencoding():
    os.reload(sys)
    sys.setdefaultencoding(default_encoding)

# Define Macros
HEADER_FILE = "template/header.docx"
FOOTER_FILE = "template/footer.docx"
DATE = str(date.today().year) + "." + str(date.today().month) + "." + str(date.today().day)


def get_doc_fullname(word_doc):
    mark = 0
    for i in range(1, len(word_doc) + 1):
        if word_doc[-i] == "\\":
            mark = i
            break
    name = word_doc[(len(word_doc) - mark + 1):]
    return name

def get_doc_name(word_doc):
    for i in range(1, len(word_doc) + 1):
        if word_doc[-i] == ".":
            mark = i
            break
    name = word_doc[(len(word_doc) - mark + 1):]
    return name

def formatting(word_doc):
    doc = Document(word_doc)
    for paragraph in doc.paragraphs:
        paragraph.add_run("\n</p>")
        paragraph.insert_paragraph_before("<p>")
    doc.save("output/" + get_doc_fullname(word_doc))

def page_build(word_doc, doc_mod, title):
    formatting(word_doc)

    doc = Document((doc_mod))
    header = Document(HEADER_FILE)
    footer = Document(FOOTER_FILE)

    add_title(header, title)

    for paragraph in doc.paragraphs:
        header.add_paragraph(paragraph.text)

    add_date(header)

    for paragraph in footer.paragraphs:
        header.add_paragraph(paragraph.text)

    header.save("output/" + get_doc_fullname(word_doc))

def add_title(doc, title):
    doc.add_paragraph("<h2>\n" + title + "\n</h2>")

def add_date(doc):
    doc.add_paragraph("<p class=\"date\">\n" + DATE + "\n</p>")


def html_maker(raw_file):
    with open("output/output.html", 'wb+') as f:
        doc = Document(raw_file)
        for paragraph in doc.paragraphs:
            f.write(paragraph.text.encode('utf8'))
        f.close()